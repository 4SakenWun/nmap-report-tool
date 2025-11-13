"""
Report generation module for scan results.
Supports PDF, DOCX, TXT, HTML, Markdown, and JSON output formats.
"""

from datetime import datetime
from typing import Dict, Optional
import os
import json


class ReportGenerator:
    """Generate professional scan reports in multiple formats."""
    
    def __init__(self, scan_results: Dict):
        self.scan_results = scan_results
        self.report_date = datetime.now()

    # ---------- Helpers ----------
    def _severity_counts(self) -> Dict[str, int]:
        counts = {"info": 0, "low": 0, "medium": 0, "high": 0, "critical": 0}
        for host in self.scan_results.get('hosts', []):
            for port in host.get('ports', []):
                sev = (port.get('severity') or 'info').lower()
                if sev in counts:
                    counts[sev] += 1
        return counts

    def _severity_color(self, sev: str) -> str:
        sev = (sev or 'info').lower()
        return {
            'critical': '#8B0000',
            'high': '#E74C3C',
            'medium': '#F39C12',
            'low': '#27AE60',
            'info': '#7F8C8D'
        }.get(sev, '#7F8C8D')

    def _severity_bar_flowable(self):
        try:
            from reportlab.graphics.shapes import Drawing, Rect
            from reportlab.lib import colors as rlcolors
        except Exception:
            return None
        counts = self._severity_counts()
        total = sum(counts.values()) or 1
        width = 400
        height = 14
        x = 0
        drawing = Drawing(width, height)
        for key in ['critical', 'high', 'medium', 'low', 'info']:
            w = int(width * (counts.get(key, 0) / total))
            if w <= 0:
                continue
            drawing.add(Rect(x, 0, w, height, fillColor=rlcolors.HexColor(self._severity_color(key))))
            x += w
        return drawing
    
    def generate_pdf(self, output_path: str) -> str:
        """
        Generate PDF report from scan results.
        
        Args:
            output_path: Path where PDF should be saved
            
        Returns:
            Path to generated PDF file
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
        except ImportError:
            raise ImportError("reportlab is required for PDF generation. Install with: pip install reportlab")
        
        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title
        story.append(Paragraph("Vulnerability Scan Report", title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Summary section
        story.append(Paragraph("Scan Summary", heading_style))
        summary_data = [
            ['Target:', self.scan_results.get('target', 'N/A')],
            ['Scan Date:', self.scan_results.get('scan_time', 'N/A')],
            ['Report Generated:', self.report_date.strftime('%Y-%m-%d %H:%M:%S')],
            ['Scanner Version:', self.scan_results.get('scanner_version', 'N/A')],
            ['Total Risk Score:', str(self.scan_results.get('total_risk', 0))]
        ]
        
        summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ecf0f1')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.2 * inch))

        # Severity summary
        sev_counts = self._severity_counts()
        sev_data = [['Severity', 'Count']]
        for k in ['critical', 'high', 'medium', 'low', 'info']:
            sev_data.append([k.capitalize(), str(sev_counts.get(k, 0))])
        sev_table = Table(sev_data, colWidths=[2*inch, 1*inch])
        sev_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(sev_table)
        bar = self._severity_bar_flowable()
        if bar:
            story.append(Spacer(1, 0.1 * inch))
            story.append(bar)
        story.append(Spacer(1, 0.3 * inch))
        
        # Host details
        for host in self.scan_results.get('hosts', []):
            story.append(Paragraph(f"Host Details", heading_style))
            
            # Host info
            addresses = ', '.join([addr['address'] for addr in host.get('addresses', [])])
            hostnames = ', '.join(host.get('hostnames', [])) or 'N/A'
            
            host_data = [
                ['IP Address(es):', addresses],
                ['Hostname(s):', hostnames],
                ['Status:', host.get('status', 'Unknown')]
            ]
            
            if host.get('os'):
                host_data.append(['OS Detection:', f"{host['os']['name']} ({host['os']['accuracy']}% accuracy)"])
            
            host_table = Table(host_data, colWidths=[2*inch, 4*inch])
            host_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ecf0f1')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            story.append(host_table)
            story.append(Spacer(1, 0.2 * inch))
            
            # Open ports
            if host.get('ports'):
                story.append(Paragraph("Open Ports and Services", heading_style))
                
                port_data = [['Port', 'Protocol', 'Service', 'Version', 'Severity']]
                for port in host.get('ports', []):
                    service = port.get('service', {})
                    service_name = service.get('name', 'unknown') if service else 'unknown'
                    version_info = ''
                    if service:
                        product = service.get('product', '')
                        version = service.get('version', '')
                        if product:
                            version_info = f"{product} {version}".strip()
                    
                    port_data.append([
                        port.get('port', ''),
                        port.get('protocol', ''),
                        service_name,
                        version_info,
                        (port.get('severity') or 'info').upper()
                    ])
                
                port_table = Table(port_data, colWidths=[0.8*inch, 0.9*inch, 1.8*inch, 1.8*inch, 0.9*inch])
                port_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
                ]))
                story.append(port_table)
                story.append(Spacer(1, 0.2 * inch))

                # Host risk
                host_risk = str(host.get('risk_score', 0))
                story.append(Paragraph(f"Risk Score: <b>{host_risk}</b>", styles['Normal']))
                story.append(Spacer(1, 0.2 * inch))
                
                # Vulnerability findings
                vuln_findings = []
                for port in host.get('ports', []):
                    for script in port.get('scripts', []):
                        if 'vuln' in script.get('id', '').lower() or script.get('output'):
                            vuln_findings.append({
                                'port': port.get('port'),
                                'script': script.get('id'),
                                'output': script.get('output', '')
                            })
                
                if vuln_findings:
                    story.append(Paragraph("Potential Vulnerabilities", heading_style))
                    for vuln in vuln_findings:
                        story.append(Paragraph(
                            f"<b>Port {vuln['port']} - {vuln['script']}</b>", 
                            styles['Normal']
                        ))
                        story.append(Paragraph(
                            vuln['output'].replace('\n', '<br/>'), 
                            styles['Normal']
                        ))
                        story.append(Spacer(1, 0.1 * inch))
        
        # Build PDF
        doc.build(story)
        return output_path

    def generate_json(self, output_path: str) -> str:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.scan_results, f, indent=2)
        return output_path

    def generate_markdown(self, output_path: str) -> str:
        lines = []
        lines.append(f"# Vulnerability Scan Report\n")
        lines.append(f"**Target:** {self.scan_results.get('target','N/A')}  ")
        lines.append(f"**Scan Date:** {self.scan_results.get('scan_time','N/A')}  ")
        lines.append(f"**Report Generated:** {self.report_date.strftime('%Y-%m-%d %H:%M:%S')}  ")
        lines.append(f"**Scanner Version:** {self.scan_results.get('scanner_version','N/A')}  ")
        lines.append(f"**Total Risk Score:** {self.scan_results.get('total_risk',0)}\n")

        sev = self._severity_counts()
        lines.append("## Severity Summary")
        lines.append("| Severity | Count |\n|---|---|")
        for k in ['critical','high','medium','low','info']:
            lines.append(f"| {k.capitalize()} | {sev.get(k,0)} |")

        # ASCII severity bar
        total = sum(sev.values()) or 1
        def bar(n):
            width = 30
            filled = int(width * (n/total))
            return '█'*filled + ' '*(width - filled)
        lines.append("\n```")
        lines.append(f"CRIT |{bar(sev['critical'])}|")
        lines.append(f"HIGH |{bar(sev['high'])}|")
        lines.append(f"MED  |{bar(sev['medium'])}|")
        lines.append(f"LOW  |{bar(sev['low'])}|")
        lines.append(f"INFO |{bar(sev['info'])}|")
        lines.append("``\n")

        for host in self.scan_results.get('hosts', []):
            lines.append("\n## Host Details")
            addresses = ', '.join([a['address'] for a in host.get('addresses', [])])
            hostnames = ', '.join(host.get('hostnames', [])) or 'N/A'
            lines.append(f"- IP: {addresses}")
            lines.append(f"- Hostnames: {hostnames}")
            lines.append(f"- Status: {host.get('status','Unknown')}")
            if host.get('os'):
                lines.append(f"- OS: {host['os']['name']} ({host['os']['accuracy']}% accuracy)")

            if host.get('ports'):
                lines.append("\n### Open Ports and Services")
                lines.append("| Port | Proto | Service | Version | Severity |\n|---|---|---|---|---|")
                for port in host.get('ports', []):
                    service = port.get('service') or {}
                    svc_name = service.get('name','unknown')
                    version_info = (service.get('product','') + ' ' + service.get('version','')).strip()
                    lines.append(
                        f"| {port.get('port','')} | {port.get('protocol','')} | {svc_name} | {version_info} | {(port.get('severity') or 'info').upper()} |"
                    )

                # Host risk
                lines.append(f"\n**Host Risk Score:** {host.get('risk_score',0)}\n")
                # Findings
                findings = []
                for p in host.get('ports', []):
                    for sc in p.get('scripts', []):
                        if 'vuln' in (sc.get('id','').lower()) or sc.get('output'):
                            findings.append((p.get('port'), sc))
                if findings:
                    lines.append("\n### Potential Vulnerabilities")
                    for portnum, sc in findings:
                        lines.append(f"- **Port {portnum} - {sc.get('id','')}**\n")
                        lines.append(f"  \n{sc.get('output','').strip()}\n")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
        return output_path

    def generate_html(self, output_path: str) -> str:
        css = """
        body{font-family:Arial,Helvetica,sans-serif;color:#1a1a1a;margin:20px}
        h1,h2,h3{color:#2c3e50}
        table{border-collapse:collapse;width:100%;margin:10px 0}
        th,td{border:1px solid #ddd;padding:8px;font-size:14px}
        th{background:#34495e;color:#fff;text-align:left}
        tr:nth-child(even){background:#f8f9fa}
        .sev{font-weight:bold;padding:2px 6px;border-radius:4px;color:#fff;display:inline-block}
        .bar{height:14px;width:100%;background:#eee;border-radius:6px;overflow:hidden}
        .seg{display:inline-block;height:100%}
        """
        def sev_span(s: str) -> str:
            color = self._severity_color(s)
            return f"<span class='sev' style='background:{color}'>{(s or 'info').upper()}</span>"

        sev_counts = self._severity_counts()
        parts = []
        parts.append(f"<html><head><meta charset='utf-8'><style>{css}</style><title>Vulnerability Scan Report</title></head><body>")
        parts.append(f"<h1>Vulnerability Scan Report</h1>")
        parts.append(f"<p><b>Target:</b> {self.scan_results.get('target','N/A')}<br/>")
        parts.append(f"<b>Scan Date:</b> {self.scan_results.get('scan_time','N/A')}<br/>")
        parts.append(f"<b>Report Generated:</b> {self.report_date.strftime('%Y-%m-%d %H:%M:%S')}<br/>")
        parts.append(f"<b>Scanner Version:</b> {self.scan_results.get('scanner_version','N/A')}<br/>")
        parts.append(f"<b>Total Risk Score:</b> {self.scan_results.get('total_risk',0)}</p>")

        parts.append("<h2>Severity Summary</h2>")
        parts.append("<table><tr><th>Severity</th><th>Count</th></tr>")
        for k in ['critical','high','medium','low','info']:
            parts.append(f"<tr><td>{k.capitalize()}</td><td>{sev_counts.get(k,0)}</td></tr>")
        parts.append("</table>")
        total = sum(sev_counts.values()) or 1
        def seg(width, color):
            return f"<span class='seg' style='width:{width}%;background:{color}'></span>"
        parts.append("<div class='bar'>")
        for k in ['critical','high','medium','low','info']:
            pct = 100 * (sev_counts.get(k,0) / total)
            parts.append(seg(pct, self._severity_color(k)))
        parts.append("</div>")

        for host in self.scan_results.get('hosts', []):
            parts.append("<h2>Host Details</h2>")
            addresses = ', '.join([a['address'] for a in host.get('addresses', [])])
            hostnames = ', '.join(host.get('hostnames', [])) or 'N/A'
            parts.append(f"<p><b>IP:</b> {addresses}<br/><b>Hostnames:</b> {hostnames}<br/><b>Status:</b> {host.get('status','Unknown')}")
            if host.get('os'):
                parts.append(f"<br/><b>OS:</b> {host['os']['name']} ({host['os']['accuracy']}% accuracy)")
            parts.append(f"<br/><b>Risk Score:</b> {host.get('risk_score',0)}")
            parts.append("</p>")

            if host.get('ports'):
                parts.append("<h3>Open Ports and Services</h3>")
                parts.append("<table><tr><th>Port</th><th>Proto</th><th>Service</th><th>Version</th><th>Severity</th></tr>")
                for port in host.get('ports', []):
                    service = port.get('service') or {}
                    svc_name = service.get('name','unknown')
                    version_info = (service.get('product','') + ' ' + service.get('version','')).strip()
                    parts.append(
                        f"<tr><td>{port.get('port','')}</td><td>{port.get('protocol','')}</td><td>{svc_name}</td><td>{version_info}</td><td>{sev_span(port.get('severity'))}</td></tr>"
                    )
                parts.append("</table>")

                findings = []
                for p in host.get('ports', []):
                    for sc in p.get('scripts', []):
                        if 'vuln' in (sc.get('id','').lower()) or sc.get('output'):
                            findings.append((p.get('port'), sc))
                if findings:
                    parts.append("<h3>Potential Vulnerabilities</h3>")
                    for portnum, sc in findings:
                        parts.append(f"<p><b>Port {portnum} - {sc.get('id','')}</b><br/>{(sc.get('output','') or '').replace('\n','<br/>')}</p>")

        parts.append("</body></html>")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("".join(parts))
        return output_path
    
    def generate_docx(self, output_path: str) -> str:
        """
        Generate DOCX report from scan results.
        
        Args:
            output_path: Path where DOCX should be saved
            
        Returns:
            Path to generated DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Vulnerability Scan Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary section
        doc.add_heading('Scan Summary', level=1)
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_data = [
            ('Target:', self.scan_results.get('target', 'N/A')),
            ('Scan Date:', self.scan_results.get('scan_time', 'N/A')),
            ('Report Generated:', self.report_date.strftime('%Y-%m-%d %H:%M:%S')),
            ('Scanner Version:', self.scan_results.get('scanner_version', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = summary_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Host details
        for host in self.scan_results.get('hosts', []):
            doc.add_heading('Host Details', level=1)
            
            addresses = ', '.join([addr['address'] for addr in host.get('addresses', [])])
            hostnames = ', '.join(host.get('hostnames', [])) or 'N/A'
            
            host_info = doc.add_table(rows=3, cols=2)
            host_info.style = 'Light Grid Accent 1'
            
            host_data = [
                ('IP Address(es):', addresses),
                ('Hostname(s):', hostnames),
                ('Status:', host.get('status', 'Unknown'))
            ]
            
            if host.get('os'):
                host_info.add_row()
                host_data.append(('OS Detection:', f"{host['os']['name']} ({host['os']['accuracy']}% accuracy)"))
            
            for i, (label, value) in enumerate(host_data):
                row = host_info.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
            
            # Open ports
            if host.get('ports'):
                doc.add_heading('Open Ports and Services', level=2)
                
                port_table = doc.add_table(rows=1, cols=4)
                port_table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = port_table.rows[0].cells
                headers = ['Port', 'Protocol', 'Service', 'Version']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                
                # Data rows
                for port in host.get('ports', []):
                    row = port_table.add_row().cells
                    service = port.get('service', {})
                    service_name = service.get('name', 'unknown') if service else 'unknown'
                    version_info = ''
                    if service:
                        product = service.get('product', '')
                        version = service.get('version', '')
                        if product:
                            version_info = f"{product} {version}".strip()
                    
                    row[0].text = str(port.get('port', ''))
                    row[1].text = port.get('protocol', '')
                    row[2].text = service_name
                    row[3].text = version_info
                
                doc.add_paragraph()
                
                # Vulnerability findings
                vuln_findings = []
                for port in host.get('ports', []):
                    for script in port.get('scripts', []):
                        if 'vuln' in script.get('id', '').lower() or script.get('output'):
                            vuln_findings.append({
                                'port': port.get('port'),
                                'script': script.get('id'),
                                'output': script.get('output', '')
                            })
                
                if vuln_findings:
                    doc.add_heading('Potential Vulnerabilities', level=2)
                    for vuln in vuln_findings:
                        p = doc.add_paragraph()
                        p.add_run(f"Port {vuln['port']} - {vuln['script']}").bold = True
                        doc.add_paragraph(vuln['output'])
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def generate_text(self, output_path: str) -> str:
        """
        Generate simple text report from scan results.
        
        Args:
            output_path: Path where text file should be saved
            
        Returns:
            Path to generated text file
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("=" * 60 + "\n")
            f.write("VULNERABILITY SCAN REPORT\n")
            f.write("=" * 60 + "\n\n")
            
            f.write("SCAN SUMMARY\n")
            f.write("-" * 60 + "\n")
            f.write(f"Target: {self.scan_results.get('target', 'N/A')}\n")
            f.write(f"Scan Date: {self.scan_results.get('scan_time', 'N/A')}\n")
            f.write(f"Report Generated: {self.report_date.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Scanner Version: {self.scan_results.get('scanner_version', 'N/A')}\n\n")
            
            for host in self.scan_results.get('hosts', []):
                f.write("HOST DETAILS\n")
                f.write("-" * 60 + "\n")
                
                addresses = ', '.join([addr['address'] for addr in host.get('addresses', [])])
                hostnames = ', '.join(host.get('hostnames', [])) or 'N/A'
                
                f.write(f"IP Address(es): {addresses}\n")
                f.write(f"Hostname(s): {hostnames}\n")
                f.write(f"Status: {host.get('status', 'Unknown')}\n")
                
                if host.get('os'):
                    f.write(f"OS Detection: {host['os']['name']} ({host['os']['accuracy']}% accuracy)\n")
                
                f.write("\n")
                
                if host.get('ports'):
                    f.write("OPEN PORTS AND SERVICES\n")
                    f.write("-" * 60 + "\n")
                    
                    for port in host.get('ports', []):
                        service = port.get('service', {})
                        service_name = service.get('name', 'unknown') if service else 'unknown'
                        version_info = ''
                        if service:
                            product = service.get('product', '')
                            version = service.get('version', '')
                            if product:
                                version_info = f"{product} {version}".strip()
                        
                        f.write(f"Port {port.get('port')}/{port.get('protocol')}: "
                               f"{service_name} {version_info}\n")
                    
                    f.write("\n")
                    
                    # Vulnerabilities
                    vuln_findings = []
                    for port in host.get('ports', []):
                        for script in port.get('scripts', []):
                            if 'vuln' in script.get('id', '').lower() or script.get('output'):
                                vuln_findings.append({
                                    'port': port.get('port'),
                                    'script': script.get('id'),
                                    'output': script.get('output', '')
                                })
                    
                    if vuln_findings:
                        f.write("POTENTIAL VULNERABILITIES\n")
                        f.write("-" * 60 + "\n")
                        for vuln in vuln_findings:
                            f.write(f"Port {vuln['port']} - {vuln['script']}\n")
                            f.write(f"{vuln['output']}\n\n")
        
        return output_path
